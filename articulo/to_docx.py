#!/usr/bin/env python3
"""Convierte manuscrito.md a un .docx con el formato de la Revista Guillermo de Ockham:
Times New Roman 12, interlineado 1,5, tamaño carta, márgenes 2,5 cm (sup/inf) y 3,0 cm (izq/der).
Tablas editables, figura insertada. python-docx puro."""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

import sys
SRC = sys.argv[1] if len(sys.argv)>1 else "/root/alejo-psico/articulo/manuscrito.md"
OUT = sys.argv[2] if len(sys.argv)>2 else "/root/alejo-psico/articulo/CliniA_manuscrito.docx"
FONT = "Times New Roman"

doc = Document()
# Estilo base Normal
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(6)

# Página carta + márgenes
sec = doc.sections[0]
sec.page_width = Cm(21.59); sec.page_height = Cm(27.94)
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0); sec.right_margin = Cm(3.0)

def style_heading(p, size, bold=True, color=None, align=None):
    for r in p.runs:
        r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold
        if color: r.font.color.rgb = color
    if align is not None: p.alignment = align

INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[\[.+?\]\])')
def add_runs(p, text):
    # divide en negritas/cursivas
    for tok in INLINE.split(text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"):
            r=p.add_run(tok[2:-2]); r.bold=True
        elif tok.startswith("*") and tok.endswith("*"):
            r=p.add_run(tok[1:-1]); r.italic=True
        elif tok.startswith("`") and tok.endswith("`"):
            r=p.add_run(tok[1:-1]); r.font.name="Courier New"
        else:
            p.add_run(tok)
    for r in p.runs:
        if r.font.name not in ("Courier New",): r.font.name=FONT
        if not r.font.size: r.font.size=Pt(12)

def add_table(rows):
    header=[c.strip() for c in rows[0].strip("|").split("|")]
    body=[[c.strip() for c in r.strip("|").split("|")] for r in rows[2:]]
    t=doc.add_table(rows=1, cols=len(header)); t.style="Table Grid"
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(header):
        cell=t.rows[0].cells[i]; cell.paragraphs[0].text=""
        add_runs(cell.paragraphs[0], h)
        for run in cell.paragraphs[0].runs: run.bold=True
    for br in body:
        cells=t.add_row().cells
        for i,val in enumerate(br):
            if i<len(cells):
                cells[i].paragraphs[0].text=""
                add_runs(cells[i].paragraphs[0], val)
    doc.add_paragraph()

lines=open(SRC).read().split("\n")
i=0
while i < len(lines):
    ln=lines[i]
    if ln.strip()=="---" or ln.strip()=="":
        i+=1; continue
    # tabla
    if ln.lstrip().startswith("|") and i+1<len(lines) and re.match(r'\s*\|[-:\s|]+\|', lines[i+1]):
        block=[]
        while i<len(lines) and lines[i].lstrip().startswith("|"):
            block.append(lines[i]); i+=1
        add_table(block); continue
    # imagen
    m=re.match(r'!\[(.*?)\]\((.*?)\)', ln.strip())
    if m:
        path=os.path.join(os.path.dirname(SRC), m.group(2))
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(6.0))
            doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        i+=1; continue
    # encabezados
    if ln.startswith("# "):
        p=doc.add_paragraph(); add_runs(p, ln[2:]); style_heading(p,14,True,RGBColor(0,0,0),WD_ALIGN_PARAGRAPH.CENTER)
    elif ln.startswith("## "):
        p=doc.add_paragraph(); add_runs(p, ln[3:]); style_heading(p,13,True)
    elif ln.startswith("### "):
        p=doc.add_paragraph(); add_runs(p, ln[4:]); style_heading(p,12,True)
    elif ln.startswith(">"):
        p=doc.add_paragraph(); add_runs(p, ln.lstrip("> ").strip())
        for r in p.runs:
            r.italic=True; r.font.color.rgb=RGBColor(0x80,0x80,0x80); r.font.size=Pt(10)
    elif re.match(r'\s*[-*]\s+', ln):
        p=doc.add_paragraph(style="List Bullet"); add_runs(p, re.sub(r'^\s*[-*]\s+','',ln))
    else:
        p=doc.add_paragraph(); add_runs(p, ln)
    i+=1

doc.save(OUT)
import subprocess
print("Guardado:", OUT)
print("Párrafos:", len(doc.paragraphs), "| Tablas:", len(doc.tables))
